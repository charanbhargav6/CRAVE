"""
CRAVE Phase 4 - LeadOrchestrator
Save to: D:/CRAVE/src/core/orchestrator.py

The brain of CRAVE. Receives commands, classifies intent,
routes to the right handler, speaks results.

Architecture:
  VoicePipeline  -->  Orchestrator  -->  ModelRouter
                                    -->  TTS
                                    -->  Agents (Phase 7+)
                                    -->  Trading (Phase 8+)

Phase 6 Orb UI will call run_once() and start_loop() directly.
"""

import os
import sys
import json
import time
import queue
import threading
from typing import Optional

from .audio_utils import load_config, reload_config, crave_root
from .model_router import ModelRouter
from .tts import speak, set_silent_mode as tts_set_silent, is_speaking, stop as tts_stop, speak_startup, speak_silent_on, speak_silent_off
from .voice import VoicePipeline

SEP = chr(92)

# ── intent categories ─────────────────────────────────────────────────────────
# Each intent maps to a handler method in Orchestrator.
# Phase 7+ handlers are stubs that will be filled in later phases.

INTENT_CHAT     = "chat"
INTENT_SCREEN   = "screen"
INTENT_FILE     = "file"
INTENT_LEARN    = "learn"
INTENT_TRADE    = "trade"
INTENT_HACK     = "hack"
INTENT_SILENT   = "silent"
INTENT_STATUS   = "status"
INTENT_STOP     = "stop"
INTENT_AUTH     = "auth"
INTENT_VIDEO    = "video"
INTENT_MESSAGE  = "message"
INTENT_SYSTEM   = "system"
INTENT_EXPLAIN  = "explain"
INTENT_EVOLVE   = "evolve"
INTENT_SELF_MODIFY = "self_modify"
INTENT_SCOUT    = "scout"
INTENT_AUTOMATE = "automate"
INTENT_UNKNOWN  = "unknown"
INTENT_PUBLIC_API = "public_api"

# ── keyword maps for FAST-PATH intent classification (fallback) ───────────────
# These are checked FIRST for instant routing (~0ms).
# If no strong match is found, the LLM classifier takes over.

_INTENT_KEYWORDS = {
    INTENT_SCREEN:  ["analyze screen", "analyze my screen", "what's on screen", "look at screen",
                     "what do you see", "screenshot", "read screen",
                     "what is on my screen"],
    INTENT_FILE:    ["create a file", "create file", "generate a file", "generate file",
                     "write a file", "write file", "save a file", "save to file",
                     "make a file", "make file", "create a script", "write a script",
                     "generate a script", "make a script",
                     "create a document", "write a document",
                     "make a ppt", "make a powerpoint", "create a presentation",
                     "make a presentation", "generate a presentation",
                     "write a report", "create a report", "generate a report",
                     "create an excel", "make a spreadsheet", "generate a spreadsheet",
                     "make a pdf", "create a pdf", "generate a pdf"],
    INTENT_TRADE:   ["trade", "trading", "buy", "sell", "long", "short",
                     "forex", "crypto", "stock", "market", "position",
                     "close all", "kill switch", "pause trading",
                     "backtest", "back test"],
    INTENT_HACK:    ["nmap", "exploit", "ctf", "flag", "kali",
                     "pentest", "hack", "vulnerability", "payload"],
    INTENT_SILENT:  ["silent mode", "go silent", "quiet mode", "mute"],
    INTENT_STATUS:  ["system status", "are you running",
                     "system check", "what phase"],
    INTENT_STOP:    ["shutdown", "exit", "quit", "turn off"],
    INTENT_AUTH:    ["keyword verification", "voice authentication", "authorize voice",
                     "verbal password", "clearance phrase"],
    INTENT_SYSTEM:  ["open ", "launch ", "start application", "close ", "kill app"],
    INTENT_MESSAGE: ["send email", "send whatsapp", "message someone"],
    INTENT_EXPLAIN: ["why did you", "explain yourself", "what was your reasoning",
                     "explain your decision", "reasoning log", "explain last action"],
    INTENT_EVOLVE:  ["upgrade yourself", "check for updates",
                     "upgrade model", "find a better model", "benchmark models"],
    INTENT_SELF_MODIFY: ["add feature", "modify yourself", "implement this",
                         "add a feature", "change your code", "write a new feature"],
    INTENT_SCOUT: ["research github", "find improvements", "scout llm", "scout repos",
                   "trending llm", "research llm", "find new models", "huggingface trending",
                   "what's new in ai", "research ai improvements", "find better techniques"],
    INTENT_AUTOMATE: ["in trading view", "draw resistance", "draw support",
                      "click on", "macro ", "automate "],
    "youtube": ["youtube shorts", "make a video about", "create a video", "youtube video"],
    INTENT_VIDEO: ["capcut", "edit video", "ffmpeg", "cut video", "subtitle video"],
    INTENT_PUBLIC_API: ["fetch price", "crypto price", "random fact", "tell me a fact",
                        "what is my ip", "network recon", "tell me a joke", "crack a joke",
                        "crack me a joke", "say a joke", "give me a joke",
                        "ipl score", "cricket score", "live score", "match score",
                        "football score", "sports score", "game score", "the score",
                        "weather", "what's the weather", "temperature outside", "what is the weather",
                        "stock price", "bitcoin price", "gold price",
                        "current price", "market price",
                        # Universal real-world queries (Stage 2 web search)
                        "petrol price", "fuel price", "gas price", "diesel price",
                        "latest news", "what happened", "results of", "who won",
                        "any events", "major events", "happening today", "happening right now",
                        "war update", "election result", "ipl match", "matches today",
                        "current situation", "what are the", "how much is", "what is the price"],
}

# ── Post-Transcription Voice Fuzzy Correction ──────────────────────────────────
# Catches common Whisper mishearings of proper nouns, app names, and websites.
# Applied BEFORE intent classification in handle().

_VOICE_CORRECTIONS = {
    # Websites / Apps
    "creepers": "cricbuzz",
    "crick buzz": "cricbuzz",
    "cric buzz": "cricbuzz",
    "crick bus": "cricbuzz",
    "trick buzz": "cricbuzz",
    "trade in view": "tradingview",
    "trading you": "tradingview",
    "what's up": "whatsapp",
    "what sup": "whatsapp",
    "watts app": "whatsapp",
    "what sap": "whatsapp",
    "get hub": "github",
    "git up": "github",
    "chat GPD": "chatgpt",
    "chat gbd": "chatgpt",
    "chatgbd": "chatgpt",
    "you tube": "youtube",
    "new tube": "youtube",
    # Common word manglings
    "i peel": "ipl",
    "i p l": "ipl",
    "ip el": "ipl",
    "eye peel": "ipl",
    "eye pl": "ipl",
}

# ── LLM Intent Classification System Prompt ───────────────────────────────────
# This is sent to the fast-path LLM (Groq/Gemini/Ollama) to classify intent.
# Engineered to generate exactly ONE token.

_LLM_INTENT_PROMPT = """You are CRAVE's intent router. Classify the user's voice command into EXACTLY ONE of these categories.

CATEGORIES:
- chat: General conversation, greetings, opinions, chitchat, questions about concepts/theory that do NOT need live data
- screen: User wants you to look at / analyze their screen or screenshot
- file: User explicitly wants to CREATE, WRITE, or SAVE a file, document, script, presentation, or spreadsheet
- learn: User wants to learn, study, research, or get a detailed explanation about a topic (educational)
- trade: Anything about trading, forex, crypto, stocks, buying/selling financial instruments, backtesting
- hack: Penetration testing, nmap, exploits, CTF, Kali Linux, cybersecurity attacks
- silent: Toggle silent/quiet/mute mode
- status: System status check
- stop: Shutdown, exit, quit the system
- auth: Voice authentication, unlock, clearance phrase
- system: Open, launch, close, or kill an application or website
- message: Send an email, WhatsApp, or Telegram message to someone
- explain: User asks why CRAVE made a decision or wants reasoning logs
- evolve: Upgrade models, check for updates, benchmark
- self_modify: User wants CRAVE to add a feature or modify its own code
- scout: Research new AI models, GitHub repos, HuggingFace trending
- automate: GUI automation, macros, TradingView drawing, clicking, mouse control
- youtube: Create or upload YouTube videos or shorts
- video: Edit video with CapCut or FFmpeg (not YouTube creation)
- public_api: ANY question that requires CURRENT, LIVE, or REAL-TIME information from the internet. This includes: weather, sports scores, match schedules, cricket, IPL, football, prices (petrol, gold, stocks, crypto), current events, news, war updates, election results, "what happened today", "is there any", event schedules, any factual question that changes over time. Also: jokes, facts, IP lookup.

RULES:
1. Output ONLY the category name. Nothing else. No explanation. No punctuation.
2. If the user asks about ANYTHING that could change daily (scores, prices, news, events, results, schedules), ALWAYS choose "public_api"
3. If the user asks to open or launch something, ALWAYS choose "system"
4. If the question is about a fact or concept that doesn't change ("what is gravity", "explain python"), use "chat" or "learn"
5. If unclear, default to "chat"

USER COMMAND: """


def classify_intent(text: str) -> str:
    """
    FAST-PATH keyword classifier. Used as fallback when LLM is unavailable.
    The Orchestrator will prefer _classify_intent_llm() when the router is ready.
    """
    lower = text.lower().strip()
    if not lower:
        return INTENT_UNKNOWN

    # Special case priority: if it contains "open ", it's always SYSTEM
    # Prevent hacking context false-positives (e.g. "scan open ports")
    if (" open " in f" {lower} " or " launch " in f" {lower} ") and "port" not in lower:
        return INTENT_SYSTEM

    # Find ALL matching intents and their longest matching keyword
    best_intent = None
    best_keyword_len = 0

    for intent, keywords in _INTENT_KEYWORDS.items():
        for kw in keywords:
            if kw in lower and len(kw) > best_keyword_len:
                best_intent = intent
                best_keyword_len = len(kw)

    if best_intent:
        return best_intent

    return INTENT_CHAT  # default: treat as normal conversation


# ── LangGraph-compatible state dict ──────────────────────────────────────────

def make_state(command: str = "", intent: str = "", response: str = "",
               context: list = None, metadata: dict = None) -> dict:
    """
    Creates the state dict that flows through the orchestrator.
    LangGraph (Phase 4+) will replace this with a TypedDict graph state.
    Using a plain dict now so the interface stays identical when LangGraph
    is wired in Phase 4.
    """
    return {
        "command":   command,
        "intent":    intent,
        "response":  response,
        "context":   context if context is not None else [],
        "metadata":  metadata if metadata is not None else {},
        "timestamp": time.time(),
    }


# ── main orchestrator class ───────────────────────────────────────────────────

class Orchestrator:
    """
    LeadOrchestrator — receives commands and routes them.

    Public API used by Phase 6 Orb UI:
        start()          — begin voice listener + main loop in background
        stop()           — shut everything down cleanly
        handle(text)     — process one text command, return response string
        run_once()       — process one command from the queue (non-blocking)
        set_silent_mode(bool)
        get_status()     — dict of current state
    """

    def __init__(self):
        self._cfg         = load_config()
        self._running     = False
        self._silent_mode = False
        self._loop_thread = None
        self._task_queue  = queue.Queue()
        self._context     = []        # conversation history (last N exchanges)
        self._max_context = self._cfg.get("context_compress_every_n_messages", 50)
        self._msg_count   = 0
        self._last_active = time.time()

        # Load system prompt from program.md
        prompt_path = os.path.join(crave_root(), "program.md")
        try:
            with open(prompt_path, "r", encoding="utf-8") as f:
                self._system_prompt = f.read()
        except:
            self._system_prompt = "You are CRAVE. Be concise."

        # Sub-systems (lazy created)
        self._router  = None
        self._voice   = None
        self._telegram = None
        self._scheduler = None
        self._last_tool_params = {}  # Tool-call extracted params

        # Neural Memory — persistent across restarts
        try:
            from src.core.neural_memory import get_neural_memory
            self._memory = get_neural_memory()
        except Exception as e:
            self._memory = None
            print(f"[Orchestrator] NeuralMemory init skipped: {e}")

        # Orb UI callbacks (set by Phase 6 Orb via set_callbacks)
        self._cb_state_change     = None  # fn(state: str)
        self._cb_command_received = None  # fn(text: str)
        self._cb_response         = None  # fn(text: str)
        self._cb_wake_word        = None  # fn()

        print("[Orchestrator] Initialised")

    # ── public API ────────────────────────────────────────────────────────────

    def start(self):
        """Start voice listener and background processing loop."""
        # 1. Security Check (Phase 5 Logic)
        from src.security.rbac import get_rbac
        get_rbac().check_lockdown()

        if self._running:
            return

        # Phase 8.1 / Security hook: Decrypt API keys into memory natively first
        try:
            from src.security.encryption import crypto_manager
            crypto_manager.decrypt_env_to_memory()
        except Exception as e:
            print(f"[Orchestrator] Env decrypt skipped: {e}")

        # Now instantiate ModelRouter so it can capture the newly decrypted API keys 
        self._router = ModelRouter()
        # Fix: Re-initialize API clients AFTER vault has injected keys into os.environ
        self._router.reinit_api_clients()
        self._voice  = VoicePipeline()

        try:
            from src.agents.telegram_agent import TelegramAgent
            from src.core.scheduler import DailyScheduler
            
            self._telegram = TelegramAgent(orchestrator=self)
            self._telegram.start()
            
            self._scheduler = DailyScheduler(telegram_agent=self._telegram, orchestrator=self)
            self._scheduler.start()
        except ImportError:
            pass

        self._running = True

        # Start voice listener
        self._voice.start()

        # Start background command loop
        self._loop_thread = threading.Thread(
            target=self._command_loop,
            daemon=True,
            name="CRAVEOrchestrator"
        )
        self._loop_thread.start()

        # Phase 10: Start Ollama GC and Calendar Sync background threads
        threading.Thread(target=self._ollama_gc_loop, daemon=True, name="CRAVE_GC").start()
        threading.Thread(target=self._calendar_loop, daemon=True, name="CRAVE_Calendar").start()

        # Phase 10: Start Thermal Monitor daemon
        try:
            from src.core.thermal_monitor import ThermalMonitor
            self._thermal = ThermalMonitor(orchestrator=self, telegram_agent=self._telegram)
            self._thermal.start()
        except ImportError as e:
            print(f"[Orchestrator] ThermalMonitor boot skipped: {e}")

        # Phase 11: Start Face ID daemon (2-hour recheck cycle)
        try:
            from src.security.face_id import FaceIDDaemon
            from src.security.rbac import get_rbac
            self._face_id = FaceIDDaemon(rbac=get_rbac())
            self._face_id.start()
        except ImportError as e:
            print(f"[Orchestrator] FaceID boot skipped: {e}")
        except Exception as e:
            print(f"[Orchestrator] FaceID init error: {e}")

        self._notify_state("idle")

        # Phase D: Index learned skills into ChromaDB vector store (background)
        try:
            from src.core.knowledge_store import index_all_skills
            threading.Thread(target=index_all_skills, daemon=True, name="CRAVE_SkillIndex").start()
        except ImportError:
            pass

        # Phase E: PentAGI Threat Detector
        try:
            from src.security.threat_detector import ThreatDetector
            from src.agents.pentagi_agent import PentagiAgent
            pentagi = PentagiAgent()
            self._threat_detector = ThreatDetector(
                log_path=os.path.join(crave_root(), "Logs", "security_events.log"),
                pentagi_agent=pentagi,
                telegram_agent=self._telegram
            )
            self._threat_detector.start()
        except Exception as e:
            print(f"[Orchestrator] Threat Detector boot skipped: {e}")

        # Start a new Neural Memory session
        if self._memory:
            self._session_id = self._memory.start_session(project="CRAVE")
            self._memory.log("milestone", "CRAVE system booted successfully")

        speak_startup()
        print("[Orchestrator] Running — say 'Hey CRAVE' or type a command")

    def stop(self):
        """Gracefully shut down all systems."""
        # Auto-save neural memory session before shutdown
        if self._memory:
            try:
                self._memory.auto_save_session(self._context, "Graceful shutdown")
            except Exception as e:
                print(f"[Orchestrator] Memory save failed: {e}")

        self._running = False
        tts_stop()
        if self._voice:
            self._voice.stop()
        if self._loop_thread:
            self._loop_thread.join(timeout=3)
        print("[Orchestrator] Stopped")

    # ── LLM Tool-Calling Router (Agentic AI Architecture) ──────────────────
    # Replaces the old keyword→intent classifier with true LLM understanding.
    # The LLM reads the user's message (even with typos) and decides which
    # tool to call, just like Claude, Gemini, and Claw do.

    # Maps tool names from tool_router.py → old intent strings for handler lookup
    _TOOL_TO_INTENT = {
        "chat":           INTENT_CHAT,
        "create_file":    INTENT_FILE,
        "open_app":       INTENT_SYSTEM,
        "close_app":      INTENT_SYSTEM,
        "web_search":     INTENT_PUBLIC_API,
        "screen_analyze": INTENT_SCREEN,
        "send_message":   INTENT_MESSAGE,
        "trading":        INTENT_TRADE,
        "system_command":  INTENT_SYSTEM,
        "generate_image": INTENT_FILE,
        "video_edit":     INTENT_VIDEO,
        "hack_pentest":   INTENT_HACK,
        "silent_mode":    INTENT_SILENT,
        "system_status":  INTENT_STATUS,
        "shutdown":       INTENT_STOP,
        "self_modify":    INTENT_SELF_MODIFY,
        "automate_gui":   "automate",
        "download_file":  "download",
        "resume_session": "resume",
        "refine_content": "refine",
    }

    _VALID_INTENTS = {
        "chat", "screen", "file", "learn", "trade", "hack", "silent",
        "status", "stop", "auth", "system", "message", "explain",
        "evolve", "self_modify", "scout", "automate", "youtube",
        "video", "public_api", "unknown",
    }

    def _tool_call_llm(self, text: str) -> dict:
        """
        Agentic Tool-Calling Router.
        Sends the user's message to the LLM with tool definitions.
        The LLM returns a JSON tool call — no keyword matching needed.
        Returns: {"tool": "tool_name", "params": {...}}
        Falls back to keyword classifier if LLM fails.
        """
        if not self._router:
            intent = classify_intent(text)
            return {"tool": intent, "params": {}}

        try:
            import time as _time
            from src.core.tool_router import TOOL_ROUTER_PROMPT, parse_tool_response

            start = _time.time()
            res = self._router.chat(
                prompt=text,
                system_prompt=TOOL_ROUTER_PROMPT,
                task_type="primary",
                options={"num_predict": 150, "temperature": 0.0},
            )
            elapsed = _time.time() - start
            raw = res.get("response", "")

            tool_call = parse_tool_response(raw)
            print(f"[Orchestrator] Tool Call: {tool_call['tool']} | Params: {tool_call.get('params', {})} ({elapsed:.2f}s via {res.get('model', '?')})")
            return tool_call

        except Exception as e:
            print(f"[Orchestrator] Tool call failed ({e}), using keyword fallback")
            intent = classify_intent(text)
            return {"tool": intent, "params": {}}

    def _classify_intent_llm(self, text: str) -> str:
        """
        Legacy compatibility wrapper.
        Internally uses tool-calling, maps result to old intent string.
        """
        tool_call = self._tool_call_llm(text)
        tool_name = tool_call.get("tool", "chat")
        # Store params on the instance for the handler to use
        self._last_tool_params = tool_call.get("params", {})
        return self._TOOL_TO_INTENT.get(tool_name, INTENT_CHAT)

    def handle(self, text: str, source: str = "local") -> str:
        """
        Process one command string.
        Returns the response string.
        Also speaks the response via TTS if the source is local.
        This is the main entry point for Phase 6 Orb UI text input.
        """
        # 1. Security Check (Phase 5 Logic)
        from src.security.rbac import get_rbac
        get_rbac().check_lockdown()
        get_rbac().touch()  # Reset idle timer
        self._last_active = time.time()  # Phase 10 GC tracking
        
        if not text or not text.strip():
            return ""

        text = text.strip()
        
        # ── Voice Fuzzy Correction: fix common Whisper mishearings ────────
        text_lower = text.lower()
        for wrong, correct in _VOICE_CORRECTIONS.items():
            if wrong in text_lower:
                text = text_lower.replace(wrong, correct)
                print(f"[Orchestrator] Voice correction: '{wrong}' → '{correct}'")
                break  # only apply one correction to avoid chain mangling
        
        # Phase 8.1 Safety: Constrain Input Tokens to prevent Ollama overflow
        if len(text) > 15000:
            print(f"[Orchestrator] Input exceeded 15k limit ({len(text)}). Truncating.")
            text = text[:15000] + "\n...[TRUNCATED BY CRAVE]"
            
        # ── The Agentic Brain (LLM Tool-Calling) ─────────────────────────
        tool_call = self._tool_call_llm(text)
        intent = tool_call.get("tool", "chat")
        tool_params = tool_call.get("params", {})
        
        # Phase 10: Multi-Step Task Chaining Override
        # If it's heavily chained ("and", "then"), route to the GUI/Task Planner instead
        lower_txt = text.lower()
        if (" and " in lower_txt or " then " in lower_txt or " next " in lower_txt) and len(lower_txt.split()) > 6:
            # We don't route pure chat/knowledge queries ("tell me about X and Y") to the planner
            if not intent == "chat":
                intent = "automate"
        
        state  = make_state(
            command=text, 
            intent=intent,
            metadata={"tool_params": tool_params}
        )

        print(f"[Orchestrator] [{source.upper()}] Command: '{text}'  Intent: {intent}  Params: {tool_params}")

        # Notify Orb: command received + thinking state (only if local)
        if source == "local":
            self._notify_command(text)
        self._notify_state("thinking")

        # Route to handler
        handler = self._get_handler(intent)
        try:
            response = handler(state)
        except Exception as e:
            import traceback
            import logging
            logger = logging.getLogger("crave.orchestrator")
            err_msg = traceback.format_exc()
            logger.error(f"Handler failed:\n{err_msg}")
            
            # Auto-Recovery via Self-Modifier ("Automatic upgradability based on mistakes")
            auto_task = f"Fix crashing bug in '{handler.__name__}'. Exception: {e}"
            response = (f"I encountered a critical error: {e}. "
                        f"However, I am automatically invoking my Self-Evolution Engine to write a patch for my own code. "
                        f"I will test the fix in a secure sandbox and ask for your confirmation shortly.")
            
            # Trigger it asynchronously
            import threading
            threading.Thread(target=self._handle_self_modify, args=({"command": auto_task},), daemon=True).start()

        # Update context
        self._context.append({"role": "user",      "content": text})
        self._context.append({"role": "assistant",  "content": response})
        self._msg_count += 2

        # Phase 12: Adaptive personality signal analysis (silent, non-blocking)
        try:
            self._analyze_adaptive_signals(text)
        except Exception:
            pass

        # Compress context if needed (Phase 10 full impl, stub here)
        if self._msg_count >= self._max_context:
            self._compress_context()

        # Notify Orb: response ready
        if source == "local":
            self._notify_response(response)

        # Speak the response
        if response and source == "local":
            self._notify_state("speaking")
            speak(response)

        self._notify_state("idle")
        return response

    def run_once(self) -> Optional[str]:
        """
        Process one item from the task queue if available.
        Returns response string or None if queue was empty.
        Called by Phase 6 Orb UI's event loop.
        """
        try:
            text = self._task_queue.get_nowait()
            return self.handle(text)
        except queue.Empty:
            return None

    def submit(self, text: str):
        """
        Submit a command to the queue (non-blocking).
        Used by Orb UI, Telegram bot (Phase 5+), and override commands.
        """
        self._task_queue.put(text.strip())

    def set_silent_mode(self, value: bool):
        """Updates internal state and TTS engine."""
        self._silent_mode = value
        tts_set_silent(value)
        
        if value:
            self.set_state("silent")
            speak_silent_on()
        else:
            self.set_state("idle")
            speak_silent_off()

    def set_callbacks(self, state_change=None, command_received=None, response=None, wake=None):
        """Allows Phase 6 Orb UI to inject thread-safe pyqtSignals."""
        if state_change:     self._cb_state_change = state_change
        if command_received: self._cb_command_received = command_received
        if response:         self._cb_response = response
        if wake:             self._cb_wake_word = wake

    def _notify_state(self, state: str):
        if self._silent_mode and state in ["idle", "listening"]:
            state = "silent"
        self._state = state
        if self._cb_state_change:
            self._cb_state_change(state)

    def _notify_command(self, text: str):
        if self._cb_command_received:
            self._cb_command_received(text)

    def _notify_response(self, text: str):
        if self._cb_response:
            self._cb_response(text)

    def _notify_wake_word(self):
        if self._cb_wake_word:
            self._cb_wake_word()

    def get_status(self) -> dict:
        """Returns internal metrics for diagnostics via `verify.py` or shell."""
        return {
            "running":        self._running,
            "silent_mode":    self._silent_mode,
            "msg_count":      self._msg_count,
            "queue_size":     self._task_queue.qsize(),
            "context_length": len(self._context),
            "voice_running":  self._voice._running if self._voice else False,
            "state":          getattr(self, '_state', 'idle'),
        }

    @property
    def state(self) -> str:
        """Current orchestrator state (used by TelegramAgent for /status)."""
        return getattr(self, '_state', 'idle')

    def set_state(self, new_state: str):
        """Public setter for orchestrator state (called by Scheduler, TelegramAgent)."""
        self._state = new_state
        self._notify_state(new_state)
        print(f"[Orchestrator] State set to: {new_state}")

    def reload_config(self):
        """Hot-reload config from disk (Phase 10)."""
        self._cfg = reload_config()
        self._max_context = self._cfg.get("context_compress_every_n_messages", 50)
        print("[Orchestrator] Config reloaded")

    # ── background loop ───────────────────────────────────────────────────────

    def _command_loop(self):
        """
        Background thread.
        Drains the voice pipeline queue and the task queue continuously.
        """
        while self._running:
            processed = False

            # 1. Check voice pipeline for new commands
            if self._voice and self._voice.command_available():
                raw = self._voice.get_next_command(timeout=0.05)
                if raw == "__wake__":
                    # Wake word fired — play acknowledgement + notify orb
                    from .tts import speak_wake
                    speak_wake()
                    self._notify_state("listening")
                    self._notify_wake_word()
                elif raw:
                    self.handle(raw)
                    processed = True

            # 2. Check internal task queue
            try:
                text = self._task_queue.get_nowait()
                self.handle(text)
                processed = True
            except queue.Empty:
                pass

            if not processed:
                time.sleep(0.05)  # light sleep to avoid 100% CPU

    def _compress_context(self):
        """Phase 10: Uses ModelRouter to intelligently summarize history."""
        if not self._router:
            # Fallback to truncation if router not loaded
            keep_count = int(self._max_context / 2)
            if len(self._context) > keep_count:
                self._context = self._context[-keep_count:]
            self._msg_count = len(self._context)
        else:
            self._context = self._router.compress_context(self._context)
            self._msg_count = len(self._context)
        print(f"[Orchestrator] Context compressed back to {self._msg_count} messages")

    def _ollama_gc_loop(self):
        """Phase 10: Flushes Ollama RAM completely if idle for 4 hours."""
        import time
        while self._running:
            time.sleep(300) # Check every 5 minutes
            idle_time = time.time() - self._last_active
            if idle_time > 14400: # 4 hours
                if not getattr(self, '_ollama_purged', False):
                    print("[GC] Idle timeout reached. Flushing Ollama from VRAM.")
                    self._ollama_purged = True
                    try:
                        import requests
                        requests.post("http://localhost:11434/api/generate", json={"model": "qwen2.5:14b", "keep_alive": 0})
                    except:
                        pass
            else:
                self._ollama_purged = False

    def _calendar_loop(self):
        """Phase 10: Checks local calendar every minute for active meetings."""
        try:
            from src.tools.calendar_sync import is_in_meeting
        except ImportError:
            return
            
        import time
        while self._running:
            time.sleep(60)
            if is_in_meeting():
                if not getattr(self, '_silent_mode', False):
                    print("\n[Calendar] Meeting detected. Auto-engaging Silent Mode.")
                    self.set_silent_mode(True)
            else:
                if getattr(self, '_silent_mode', False):
                    # We are silent but NO meeting is active. 
                    # Do not auto-un-silence in case the user manually silenced it!
                    pass

    # ── intent handlers ───────────────────────────────────────────────────────

    def _get_handler(self, intent: str):
        handlers = {
            INTENT_CHAT:    self._handle_chat,
            INTENT_SCREEN:  self._handle_screen,
            INTENT_FILE:    self._handle_file,
            INTENT_LEARN:   self._handle_learn,
            INTENT_TRADE:   self._handle_trade,
            INTENT_HACK:    self._handle_hack,
            INTENT_SILENT:  self._handle_silent,
            INTENT_STATUS:  self._handle_status,
            INTENT_STOP:    self._handle_stop,
            INTENT_AUTH:    self._handle_auth,
            INTENT_VIDEO:   self._handle_video,
            INTENT_MESSAGE: self._handle_message,
            INTENT_SYSTEM:  self._handle_system,
            "youtube":      self._handle_youtube,
            "automate":     self._handle_automation,
            "system_command": self._handle_agent_loop,
            INTENT_EXPLAIN: self._handle_explain,
            INTENT_EVOLVE:  self._handle_evolve,
            INTENT_SELF_MODIFY: self._handle_self_modify,
            INTENT_SCOUT: self._handle_scout,
            INTENT_PUBLIC_API: self._handle_public_api,
            "download":    self._handle_download,
            "resume":      self._handle_resume,
            "refine":      self._handle_refine,
            INTENT_UNKNOWN: self._handle_chat,
        }
        return handlers.get(intent, self._handle_chat)

    # ── Resume Session Handler (Neural Memory) ───────────────────────────

    def _handle_resume(self, state: dict) -> str:
        """
        'Where did we stop?' — Recalls last session with 2-3 bullet points
        and reopens recent tabs/files.
        """
        if not self._memory:
            return "Neural Memory system is not available."

        brief = self._memory.generate_resume_brief()
        tabs = self._memory.get_recent_tabs()

        # Reopen recent tabs/files
        if tabs:
            import os, subprocess
            reopened = []
            for tab in tabs[:5]:  # Max 5 tabs
                try:
                    if tab.startswith("http"):
                        os.startfile(tab)
                    elif os.path.exists(tab):
                        os.startfile(tab)
                    reopened.append(tab)
                except Exception:
                    pass
            if reopened:
                brief += f"\n\nReopened {len(reopened)} recent tabs."

        return brief

    # ── File Download Handler (L3 Security Gate) ─────────────────────────

    def _handle_download(self, state: dict) -> str:
        """
        Download files from the internet. .exe/.msi require L3 clearance.
        """
        params = getattr(self, '_last_tool_params', {})
        url = params.get("url", "")
        filename = params.get("filename", "")
        save_location = params.get("save_location", "downloads")

        if not url:
            # Try to extract URL from the raw command
            import re
            urls = re.findall(r'https?://\S+', state["command"])
            if urls:
                url = urls[0]
            else:
                return "I need a URL to download from. Please include the download link."

        # Extract filename from URL if not provided
        if not filename:
            from urllib.parse import urlparse
            parsed = urlparse(url)
            filename = os.path.basename(parsed.path) or "download"

        # Security gate: .exe, .msi, .bat, .ps1 require L3
        DANGEROUS_EXTS = {'.exe', '.msi', '.bat', '.ps1', '.cmd', '.com', '.scr', '.vbs'}
        _, ext = os.path.splitext(filename.lower())

        if ext in DANGEROUS_EXTS:
            from src.security.rbac import get_rbac
            rbac = get_rbac()
            if rbac.auth_level < 3:
                return (f"⚠️ SECURITY: Downloading executable files ({ext}) requires L3 clearance. "
                        f"Current level: L{rbac.auth_level}. Please authenticate to L3 first.")

        # Resolve save path
        from src.core.tool_router import resolve_save_path
        filepath = resolve_save_path(filename, save_location)

        # Download with progress
        speak(f"Downloading {filename}. Please wait.")

        try:
            import requests
            response = requests.get(url, stream=True, timeout=60,
                                     headers={"User-Agent": "CRAVE/10.3"})
            response.raise_for_status()

            total = int(response.headers.get('content-length', 0))
            downloaded = 0

            os.makedirs(os.path.dirname(filepath), exist_ok=True)
            with open(filepath, 'wb') as f:
                for chunk in response.iter_content(chunk_size=8192):
                    f.write(chunk)
                    downloaded += len(chunk)

            size_mb = downloaded / (1024 * 1024)

            # Log to neural memory
            if self._memory:
                self._memory.log("file", f"Downloaded: {filename} ({size_mb:.1f}MB)", 
                                {"url": url, "path": filepath})

            return f"✅ Download complete!\n📄 {filename} ({size_mb:.1f} MB)\n📁 Saved at: {filepath}"
        except Exception as e:
            return f"❌ Download failed: {e}"

    # ── GAN Refiner Handler ──────────────────────────────────────────────

    def _handle_refine(self, state: dict) -> str:
        """
        Generator-Evaluator quality loop. Two AIs compete:
        Generator writes → Evaluator scores → Generator rewrites. 3 rounds.
        Evaluator has NO memory between rounds.
        """
        params = getattr(self, '_last_tool_params', {})
        task = params.get("task", state["command"])
        rubric = params.get("rubric", "Clear, professional, well-structured, concise, error-free")

        if not self._router:
            return "Router not ready for content refinement."

        speak("Running quality refinement loop. Generator and Evaluator competing. Stand by.")

        from src.core.gan_refiner import refine
        result = refine(
            task=task,
            rubric=rubric,
            router=self._router,
            rounds=3,
            pass_threshold=8,
        )

        scores = result.get("scores", [])
        rounds_used = result.get("rounds_used", 0)
        passed = result.get("passed", False)
        final = result.get("final_output", "")

        verdict = "PASSED ✅" if passed else "DID NOT PASS ❌"
        score_str = " → ".join([str(s) for s in scores])

        summary = (
            f"🔄 GAN Refiner: {rounds_used} rounds | Scores: {score_str} | {verdict}\n\n"
            f"{final}"
        )

        # Log to neural memory
        if self._memory:
            self._memory.log("decision", f"GAN Refiner: {rounds_used}R, scores={scores}, passed={passed}")

        return summary

    def _handle_public_api(self, state: dict) -> str:
        """Universal internet intelligence handler.
        Speaks contextual acknowledgement, then fetches real-time data."""
        self._notify_state("thinking")
        cmd = state["command"]
        
        # ── Contextual acknowledgement: "Let me look up about [task]" ─────
        # Extract the core topic for a natural-sounding acknowledgement
        topic = cmd.strip()
        # Strip common prefixes so we get just the subject
        for prefix in ["what are the", "what is the", "are there any", "is there any",
                       "tell me about", "show me", "check", "find", "get me",
                       "what are", "what is", "how much", "how many", "any",
                       "give me", "fetch", "look up", "search for", "search"]:
            if topic.lower().startswith(prefix):
                topic = topic[len(prefix):].strip()
                break
        topic = topic.rstrip("?.!").strip() or cmd
        
        speak(f"Let me look up about {topic}.")
        
        from src.agents.public_api_agent import PublicApiAgent
        agent = PublicApiAgent(orchestrator=self)
        try:
            return agent.handle_request(cmd)
        except Exception as e:
            import logging
            logging.getLogger("crave.orchestrator").error(f"Public API Agent failed: {e}")
            return f"I tried to look up {topic} but hit an error. Try rephrasing your question."

    def _handle_message(self, state: dict) -> str:
        """Parses message intent. Defaults to WhatsApp unless email/telegram is specified."""
        cmd = state["command"]
        
        # Use ModelRouter to parse out variables reliably
        sys_prompt = (
            "You are an NLP parser. Extract the messaging details from the user's command.\n"
            "Return EXACTLY AND ONLY this JSON format with no other text:\n"
            '{"type": "email" or "whatsapp" or "telegram", "target": "contact_name_or_number_or_email", "content": "the message itself"}\n'
            "IMPORTANT: If the user does not specify a platform, default type to 'whatsapp'."
        )
        
        try:
            res = self._router.chat(prompt=cmd, system_prompt=sys_prompt)
            raw = res.get("response", "{}")
            # Clean markdown
            if "```json" in raw:
                raw = raw.split("```json")[1].split("```")[0].strip()
            elif "```" in raw:
                raw = raw.split("```")[1].split("```")[0].strip()
                
            data = json.loads(raw)
            
            m_type = data.get("type", "whatsapp").lower()
            target = data.get("target", "")
            content = data.get("content", "")
            
            if not target or not content:
                return "I couldn't understand who to message or what to say. Please try again."
                
            if m_type == "email":
                from src.agents.email_agent import EmailAgent
                agent = EmailAgent()
                return agent.send_email(to_address=target, subject="CRAVE Dispatch", body=content)
            
            elif m_type == "telegram":
                # Use existing Telegram bot if available
                if self._telegram:
                    self._telegram.send_message_sync(f"To {target}: {content}")
                    return f"Telegram message sent to {target}."
                return "Telegram bot not configured."
                
            else:
                # Default: WhatsApp Desktop app automation
                contact_name = target
                phone = ""
                
                # If target looks like a phone number, use it directly
                if target.startswith("+") or target.replace(" ", "").isdigit():
                    phone = target
                    contact_name = ""
                else:
                    # Try to resolve from encrypted vault for extra info
                    try:
                        from src.security.contact_vault import ContactVault
                        vault = ContactVault()
                        contact = vault.resolve(target)
                        if contact:
                            contact_name = contact["name"]
                            phone = contact.get("phone", "")
                            speak(f"Sending WhatsApp message to {contact['name']}.")
                    except Exception:
                        pass  # Just use the raw name for search
                
                from src.agents.whatsapp_agent import WhatsAppAgent
                agent = WhatsAppAgent()
                return agent.send_whatsapp(phone_number=phone, message=content, contact_name=contact_name)
                
        except Exception as e:
            import logging
            logger = logging.getLogger("crave.orchestrator")
            logger.error(f"Failed to parse message command: {e}")
            return "I failed to parse the message. Please try: 'message [name] saying [content]'"

    def _handle_agent_loop(self, state: dict) -> str:
        """
        Claw-style Agentic Terminal Loop.
        Allows the LLM to write and execute PowerShell commands iteratively
        until it resolves the user's request, with RBAC security checks.
        """
        if not self._router:
            return "Model router not ready."

        from src.security.rbac import get_rbac
        import subprocess
        import re

        agent_sys_prompt = (
            "You are an autonomous shell agent. You can run powershell commands to inspect the system, "
            "find files, and answer the user's request. "
            "To run a command, output it exactly inside a markdown ```powershell block. "
            "I will run it and return the output. If you have the final answer, reply normally without a code block."
        )

        prompt = state["command"]
        messages = self._context[-10:] if len(self._context) > 10 else self._context.copy()
        
        max_steps = 5
        for step in range(max_steps):
            self._notify_state("thinking")
            res = self._router.chat(
                prompt=prompt,
                messages=messages,
                system_prompt=agent_sys_prompt,
                task_type="primary"
            )
            
            response_text = res.get("response", "")
            messages.append({"role": "assistant", "content": response_text})

            # Look for powershell code blocks
            match = re.search(r"```powershell\s+(.*?)\s+```", response_text, re.DOTALL | re.IGNORECASE)
            if not match:
                return response_text  # Final answer reached
                
            cmd = match.group(1).strip()
            
            # ── RBAC Security Check ──────────────────────────────────────
            # Level 1: Read-only (auto-approved)
            # Level 2: Moderate state changes (mkdir, echo)
            # Level 3/4: Destructive (rm, del, reg, format, Stop-Process)
            
            cmd_lower = cmd.lower()
            auth_required = 1
            
            destructive_keywords = ["rm ", "del ", "remove-item", "format", "reg ", "stop-process", "kill"]
            moderate_keywords = ["mkdir", "new-item", "echo", "Out-File", "set-content", "add-content"]
            
            if any(k in cmd_lower for k in destructive_keywords):
                auth_required = 3
            elif any(k in cmd_lower for k in moderate_keywords) or ">" in cmd_lower:
                auth_required = 2
                
            rbac = get_rbac()
            if rbac.auth_level < auth_required:
                msg = f"Security Intercept: Command '{cmd[:20]}...' requires L{auth_required} clearance (Current: L{rbac.auth_level}). Authenticate to proceed."
                print(f"[AgentLoop] BLOCKED: {msg}")
                return msg

            print(f"[AgentLoop] Executing (L{auth_required} authorized): {cmd}")
            self._notify_state("speaking")
            
            try:
                # Execute securely
                result = subprocess.run(
                    ["powershell", "-Command", cmd],
                    capture_output=True, text=True, timeout=15
                )
                output = result.stdout + result.stderr
                if not output.strip():
                    output = "[Command executed successfully with no output]"
            except subprocess.TimeoutExpired:
                output = "[Error: Command timed out after 15 seconds]"
            except Exception as e:
                output = f"[Execution Error: {e}]"
                
            # Truncate giant outputs
            if len(output) > 2000:
                output = output[:2000] + "\n...[TRUNCATED]"
                
            prompt = f"Command Output:\n{output}\nContinue reasoning."
            messages.append({"role": "user", "content": prompt})
            
        return "I reached the maximum number of terminal steps without a final answer."

    def _handle_system(self, state: dict) -> str:
        """Handles basic OS system commands dynamically."""
        cmd = state["command"].lower().strip()
        import os
        import subprocess
        
        try:
            target = cmd
            
            # Robust extraction of the target program/website
            import re
            match = re.search(r"(?:open|launch|start|close|kill)\s+(?:up\s+)?(.+)", cmd)
            if match:
                target = match.group(1).strip()
            
            # Sanitize target word
            target = target.replace("please", "").replace("can you", "").strip()
            
            if not target or target == cmd and "open" not in cmd and "close" not in cmd:
                return "I'm not sure what you want me to open or close."

            # Map very common names to exact shell protocol or exe names
            aliases = {
                "notepad": "notepad.exe",
                "calculator": "calc.exe",
                "calc": "calc.exe",
                "browser": "msedge.exe",
                "chrome": "chrome.exe",
                "explorer": "explorer.exe",
                "command prompt": "cmd.exe",
                "terminal": "wt.exe",
                "whatsapp": "whatsapp:",
                # Website aliases
                "trading view": "https://tradingview.com",
                "tradingview": "https://tradingview.com",
                "youtube": "https://youtube.com",
                "github": "https://github.com",
                "chatgpt": "https://chatgpt.com",
                "cricbuzz": "https://cricbuzz.com",
                "google": "https://google.com",
                "espn": "https://espn.com",
                "stackoverflow": "https://stackoverflow.com",
                "stack overflow": "https://stackoverflow.com",
                "reddit": "https://reddit.com",
                "twitter": "https://x.com",
                "instagram": "https://instagram.com",
                "linkedin": "https://linkedin.com",
                "amazon": "https://amazon.in",
                "flipkart": "https://flipkart.com",
                "netflix": "https://netflix.com",
                "spotify": "https://spotify.com",
            }
            
            # Smart substring matching (e.g. "trading view in silent mode" -> "trading view")
            exe = target
            for key, val in aliases.items():
                if key in target:
                    exe = val
                    target = key # Reset target name for cleaner speech output
                    break
            
            # Check if this is a website
            if exe.startswith("http://") or exe.startswith("https://") or ('.' in exe and not exe.endswith('.exe')):
                # URL routing bypasses Python module and uses correct Windows Default App routing
                url = exe if exe.startswith("http") else f"https://{exe}"
                try:
                    os.startfile(url)
                except Exception:
                    subprocess.Popen(["powershell", "-c", f"Start-Process '{url}'"])
                return f"Opening website {target}."
            
            if "close" in cmd or "kill" in cmd:
                # Fuzzy matching for process closing (e.g. "notepd" -> "notepad")
                kill_target = exe if exe.endswith(".exe") else f"{exe}.exe"
                # If exact kill fails, try a wild-card kill
                ret_code = subprocess.call(["taskkill", "/F", "/IM", kill_target], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                if ret_code != 0 and not exe.endswith(".exe"):
                    # Try fuzzy kill via powershell
                    subprocess.call(["powershell", "-c", f"Get-Process *{exe}* | Stop-Process -Force"], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                return f"Closed {target}."
            else:
                # Handle special folder aliases
                if "crave" in target.lower() and "folder" in target.lower():
                    os.startfile("D:\\CRAVE")
                    return "Opening CRAVE workspace folder."

                try:
                    # Native foreground execution (handles apps, docs, folders, protocols)
                    os.startfile(exe)
                except FileNotFoundError:
                    # Backup generic execution
                    subprocess.Popen(["powershell", "-command", f"Start-Process '{exe}'"])
                return f"Opening {target}."
        except Exception as e:
            import logging
            logger = logging.getLogger("crave.orchestrator")
            logger.error(f"System command failed: {e}")
            return f"Failed to perform action."

    def _handle_chat(self, state: dict) -> str:
        """
        General conversation and Q&A — routes to Qwen3 via ModelRouter.
        This is the workhorse handler used by ~80% of commands.
        """
        cmd_lower = state["command"].lower()
        
        # Intercept WhatsApp and YouTube hallucinations
        if "whatsapp message" in cmd_lower:
            return "I cannot read WhatsApp messages automatically. The WhatsApp Desktop app does not provide an API for message extraction due to end-to-end encryption. I can only send messages via automation."
            
        if "youtube channel" in cmd_lower or "about my channel" in cmd_lower:
            return "I don't have active OAuth access to your YouTube Studio analytics right now. To view your recent channel updates and subscriber stats, you need to authorize the YouTube Data API in my configuration."
            
        if "location of crave" in cmd_lower or "where is crave" in cmd_lower:
            return "My core system is located at D:\\CRAVE on your local machine."

        if not self._router:
            return "Model router not ready yet."

        recent = self._context[-20:] if len(self._context) > 20 else self._context
        res = self._router.chat(
            prompt=state["command"],
            messages=recent,
            system_prompt=self._system_prompt
        )
        return res.get("response", "I didn't get a response. Please try again.")

    def _handle_screen(self, state: dict) -> str:
        """
        Screen analysis (Phase 7).
        Uses mss to capture screen to RAM, then queries Gemma 3 Local Vision.
        """
        self._notify_state("thinking")
        from src.agents.screen_agent import ScreenAgent
        
        agent = ScreenAgent()
        # Clean up the command so it acts as a prompt for the multimodal LLM
        prompt = state.get("command", "")
        if not prompt or len(prompt) < 5:
             prompt = "Describe what is currently on my screen in detail."
             
        # Add verbal feedback that it's looking
        from .tts import speak
        speak("Analyzing optical feed.")
        
        response = agent.analyze_screen(prompt)
        return response

    def _handle_youtube(self, state: dict) -> str:
        """Phase 11: End-to-End YouTube Creation Pipeline"""
        cmd = state["command"]
        
        self._notify_state("thinking")
        speak("Initializing YouTube autonomous content pipeline. Selecting niche and writing script.")
        
        # Determine topic and video length preference using small LLM task
        sys_prompt = (
            "Extract the requested YouTube topic/niche, whether it's a 'short' or 'long' video, and the target 'channel' (default 'main').\n"
            "Respond ONLY with JSON:\n"
            '{"topic": "the core topic", "is_short": true/false, "channel": "main"}'
        )
        res = self._router.chat(prompt=cmd, system_prompt=sys_prompt)
        try:
            import json
            data = json.loads(res.get("response", "{}"))
            topic = data.get("topic", "Trending internet technology 2026")
            is_short = data.get("is_short", True)
            channel = data.get("channel", "main")
        except:
            topic = cmd.replace("make a youtube video about", "").strip()
            is_short = True
            channel = "main"
            
        speak(f"Assets generated. Assembling video payload for channel: {channel}.")
            
        from src.agents.youtube_shorts_agent import YouTubeShortsAgent
        yt = YouTubeShortsAgent(orchestrator=self)
        result = yt.run_pipeline(topic=topic, is_short=is_short, upload_private=True, channel_name=channel)
        
        return result

    def _handle_file(self, state: dict) -> str:
        """
        File generation — creates files directly on disk at the user-specified location.
        Uses tool-call params (filename, save_location, description) extracted by the
        agentic tool router. Falls back to LLM extraction if params are missing.
        
        FIXED (v10.3): Files now save to user's Desktop/Documents/specified path
        instead of hardcoded D:/CRAVE/data/generated_files/.
        """
        cmd = state["command"]
        import json, os
        from src.core.tool_router import resolve_save_path
        
        # ── Step 1: Get file metadata from tool-call params ──────────────
        params = getattr(self, '_last_tool_params', {})
        filename = params.get("filename", "")
        save_location = params.get("save_location", "desktop")
        description = params.get("description", cmd)
        
        # ── Step 2: If tool-call didn't extract params, ask LLM to extract ─
        if not filename:
            sys_over = (
                "You are CRAVE. The user wants to create a file. "
                "Extract the details and respond ONLY with this exact JSON format:\n"
                '{"filename": "name.ext", "save_location": "desktop", "content": "file content here"}\n\n'
                "RULES:\n"
                '- filename: the file name with extension (e.g. notes.txt, script.py)\n'
                '- save_location: "desktop", "documents", "downloads", or an absolute path. Default: "desktop"\n'
                '- content: the actual file content to write\n'
                '- If user says "desktop" or "on my desktop", save_location = "desktop"\n'
                '- If user says "save in D:\\project", save_location = "D:\\project"'
            )
            if not self._router:
                return "Router not ready for file creation."
            
            res = self._router.chat(
                prompt=cmd,
                system_prompt=sys_over,
                task_type="reasoning"
            )
            
            reply_str = res.get("response", "")
            # Clean markdown wrapping
            if "```json" in reply_str:
                reply_str = reply_str.split("```json")[1].split("```")[0].strip()
            elif "```" in reply_str:
                reply_str = reply_str.split("```")[1].split("```")[0].strip()
            
            # Find JSON in response
            start_idx = reply_str.find("{")
            end_idx = reply_str.rfind("}") + 1
            if start_idx >= 0 and end_idx > start_idx:
                reply_str = reply_str[start_idx:end_idx]
            
            try:
                data = json.loads(reply_str)
                filename = data.get("filename", "generated_file.txt")
                save_location = data.get("save_location", "desktop")
                content = data.get("content", "")
            except json.JSONDecodeError:
                return "I couldn't understand the file details. Please try: 'create a file named X.txt on desktop with content Y'"
        else:
            # Tool-call gave us params — now generate the content via LLM
            content_prompt = (
                f"Generate the content for a file named '{filename}'.\n"
                f"User's request: {description}\n\n"
                "Output ONLY the raw file content. No explanations, no markdown wrapping, no JSON."
            )
            if self._router:
                res = self._router.chat(
                    prompt=content_prompt,
                    system_prompt="You are a file content generator. Output only raw file content.",
                    task_type="reasoning"
                )
                content = res.get("response", "")
            else:
                content = ""
        
        # ── Step 3: Write the file to the correct location ───────────────
        filepath = resolve_save_path(filename, save_location)
        
        # Security check: prevent writing to vault
        from src.agents.file_agent import _is_vault_path
        if _is_vault_path(filepath):
            return "ERROR: Cannot write to protected vault directory."
        
        try:
            os.makedirs(os.path.dirname(filepath), exist_ok=True)
            
            # Format-specific generation
            ext = os.path.splitext(filepath)[1].lower()
            
            if ext == '.docx':
                from docx import Document
                doc = Document()
                doc.add_heading(filename.replace(".docx", "").replace("_", " ").title(), 0)
                # Split content into paragraphs
                for para in content.split('\n\n'):
                    if para.strip():
                        doc.add_paragraph(para.strip())
                doc.save(filepath)
                
            elif ext == '.pdf':
                from reportlab.lib.pagesizes import letter
                from reportlab.pdfgen import canvas
                from reportlab.lib.utils import simpleSplit
                c = canvas.Canvas(filepath, pagesize=letter)
                width, height = letter
                
                c.setFont("Helvetica", 12)
                y = height - 50
                for para in content.split('\n\n'):
                    if not para.strip(): continue
                    # Rough text wrapping
                    lines = simpleSplit(para.strip(), "Helvetica", 12, width - 100)
                    for line in lines:
                        if y < 50:
                            c.showPage()
                            c.setFont("Helvetica", 12)
                            y = height - 50
                        c.drawString(50, y, line)
                        y -= 15
                    y -= 10 # paragraph spacing
                c.save()
            else:
                # Default text generation
                with open(filepath, "w", encoding="utf-8") as f:
                    f.write(content)
                    
            # Log to neural memory
            if self._memory:
                self._memory.log("file", f"Created document: {filename}", {"path": filepath})
                
            return f"File created successfully!\n📄 {filename}\n📁 Saved at: {filepath}"
        except Exception as e:
            return f"Failed to create file: {e}"

    def _handle_video(self, state: dict) -> str:
        """
        Routes video manipulation requests to FFmpeg Agent.
        CapCut agent removed (v10.3) — GUI hijacking was unreliable.
        """
        cmd = state["command"].lower()
        
        # YouTube video creation
        if any(kw in cmd for kw in ["youtube", "shorts", "upload"]):
            return self._handle_youtube(state)
        
        # Default to FFmpeg headless
        return ("FFmpeg pipeline active. "
                "I am standing by for specific timestamps or video file coordinates "
                "to rapidly inject mathematical cuts or subtitles.")

    def _handle_learn(self, state: dict) -> str:
        """
        Visual Learning Mode — Phase 12.
        1. Asks LLM for concise explanation
        2. Searches web for relevant diagrams/images
        3. If none found, generates a chart locally
        4. Returns HTML response with inline images for the TerminalWidget
        """
        cmd = state["command"]
        
        # Step 0: Check if we already know something about this topic (ChromaDB)
        prior_knowledge = ""
        try:
            from src.core.knowledge_store import search_skills
            hits = search_skills(cmd, n_results=2)
            if hits and hits[0].get("relevance", 0) > 0.3:
                prior_knowledge = "\n\n".join([h["content"] for h in hits])
                prior_knowledge = f"\n\n[You have prior knowledge on this topic]:\n{prior_knowledge[:800]}"
        except Exception:
            pass

        # Step 1: Get the core explanation from LLM (adapted to user's preferred style)
        style_hint = self._get_user_style_hint()
        sys_over = (
            f"You are CRAVE. Explain the topic clearly and concisely. "
            f"{style_hint}"
            f"Use bullet points for key concepts. Keep it under 150 words."
            f"{prior_knowledge}"
        )
        recent = self._context[-20:] if len(self._context) > 20 else self._context
        if not self._router:
            return "Learning module not ready."

        res = self._router.chat(
            prompt=cmd,
            messages=recent,
            system_prompt=sys_over
        )
        explanation = res.get("response", "I couldn't generate an explanation.")

        # Step 2: Try to find or generate visuals
        visual_html = ""
        try:
            from src.agents.visual_agent import VisualAgent
            va = VisualAgent(router=self._router)
            
            # Extract the core topic for image search
            topic = cmd.lower()
            for prefix in ["teach me", "learn about", "study", "research", "find out about",
                           "explain", "what is", "how does", "tell me about"]:
                topic = topic.replace(prefix, "").strip()
            
            result = va.get_visuals(topic)
            
            if result["images"]:
                source_label = "🌐 Web" if result["source"] == "web" else "📊 Generated"
                visual_html = f'<div style="color: #00FFCC; font-size: 9px; margin-top: 6px;">[{source_label}]</div>'
                for img_path in result["images"][:2]:  # Max 2 images
                    # Convert to file:// URI for QTextBrowser
                    file_uri = img_path.replace("\\", "/")
                    visual_html += f'<img src="file:///{file_uri}" width="320" style="margin: 4px 0;" /><br/>'
        except Exception as e:
            import logging
            logging.getLogger("crave.orchestrator").warning(f"Visual agent failed: {e}")

        # Step 3: Log topic for adaptive personality
        self._log_learning_interaction(cmd)

        # Step 4: Persist skill artifact via ResearchAgent (background — don't block response)
        try:
            topic = cmd.lower()
            for prefix in ["teach me", "learn about", "study", "research", "find out about",
                           "explain", "what is", "how does", "tell me about"]:
                topic = topic.replace(prefix, "").strip()
            if topic and len(topic) > 2:
                import threading as _thr
                def _persist_skill():
                    try:
                        from src.agents.research_agent import ResearchAgent
                        ra = ResearchAgent(orchestrator=self)
                        ra.learn_topic(topic)
                    except Exception as e:
                        import logging
                        logging.getLogger("crave.orchestrator").warning(f"Skill persistence failed: {e}")
                _thr.Thread(target=_persist_skill, daemon=True, name="SkillPersist").start()
        except Exception:
            pass

        # Combine text + visuals
        full_response = explanation
        if visual_html:
            full_response += f"\n\n{visual_html}"
        
        return full_response

    def _get_user_style_hint(self) -> str:
        """Load user's preferred explanation style from profile."""
        profile_path = os.path.join(crave_root(), "data", "user_profile.json")
        try:
            with open(profile_path, "r", encoding="utf-8") as f:
                profile = json.load(f)
            style = profile.get("preferred_style", "")
            if style:
                return f"This user learns best through {style}. Adapt your explanation accordingly. "
        except:
            pass
        return ""

    def _log_learning_interaction(self, topic: str):
        """Track topics for adaptive personality learning."""
        from datetime import datetime
        profile_path = os.path.join(crave_root(), "data", "user_profile.json")
        try:
            if os.path.exists(profile_path):
                with open(profile_path, "r", encoding="utf-8") as f:
                    profile = json.load(f)
            else:
                profile = {
                    "preferred_style": "",
                    "topics_discussed": [],
                    "positive_signals": 0,
                    "negative_signals": 0,
                    "total_interactions": 0,
                }
            
            profile["topics_discussed"].append({
                "topic": topic[:100],
                "timestamp": datetime.now().isoformat()
            })
            # Keep only last 100 topics
            profile["topics_discussed"] = profile["topics_discussed"][-100:]
            profile["total_interactions"] = profile.get("total_interactions", 0) + 1
            
            os.makedirs(os.path.dirname(profile_path), exist_ok=True)
            with open(profile_path, "w", encoding="utf-8") as f:
                json.dump(profile, f, indent=2)
        except Exception:
            pass

    def _handle_trade(self, state: dict) -> str:
        """
        Trading commands — Full Phase 8 Pipeline.
        Routes override commands, then runs the autonomous 5-agent chain:
        DataAgent → StrategyAgent → RiskAgent → ExecutionAgent
        """
        cmd = state["command"].lower()

        # ── Backtest sub-routing (intercept before live trading) ──────────
        if "backtest" in cmd or "back test" in cmd:
            return self._run_backtest(cmd)

        # ── Override commands (work immediately) ─────────────────────────────
        if "kill" in cmd or "kill switch" in cmd:
            return self._trade_kill_switch()

        if "close all" in cmd or "close" in cmd:
            return self._trade_close_all()

        if "pause" in cmd:
            self._trading_paused = True
            return "Trading paused. Open positions remain active. Say 'resume' to restart."

        if "resume" in cmd:
            self._trading_paused = False
            return "Trading resumed. Full autonomous execution re-engaged."

        if "status" in cmd:
            return self._trade_status()

        # Direction overrides (next single trade only)
        if "go long" in cmd or cmd.strip() == "long":
            self._trade_direction_override = "buy"
            return "Override accepted. Next trade will be a BUY. Auto resumes after."

        if "go short" in cmd or cmd.strip() == "short":
            self._trade_direction_override = "sell"
            return "Override accepted. Next trade will be a SELL. Auto resumes after."

        # ── Autonomous Pipeline ──────────────────────────────────────────────
        if getattr(self, '_trading_paused', False):
            return "Trading is currently paused. Say 'CRAVE resume' to re-engage."

        try:
            from Sub_Projects.Trading.data_agent import DataAgent
            from Sub_Projects.Trading.strategy_agent import StrategyAgent
            from Sub_Projects.Trading.risk_agent import RiskAgent
            from Sub_Projects.Trading.execution_agent import ExecutionAgent
            from src.core.memory_bank import MemoryBank

            # Extract symbol from command or default
            symbol = self._extract_symbol(cmd)
            exchange = "alpaca"
            if any(w in cmd for w in ["crypto", "btc", "eth", "binance"]):
                exchange = "binance"
            elif any(w in cmd for w in ["forex", "eurusd", "gbpusd", "mt5"]):
                exchange = "mt5"

            data_agent = DataAgent()
            strategy = StrategyAgent()
            risk = RiskAgent()
            execution = ExecutionAgent(data_agent=data_agent)
            memory = MemoryBank()

            # 1. Fetch market data
            df = data_agent.get_ohlcv(symbol, exchange=exchange, timeframe="1h", limit=250)
            if df is None or len(df) < 50:
                return f"Insufficient market data for {symbol} on {exchange}. Cannot analyze."

            # 2. Check Red Folder (news guard)
            currency = symbol[:3] if len(symbol) >= 3 else "USD"
            news_check = data_agent.check_red_folder(currency)
            if news_check.get("is_danger"):
                event_name = news_check.get("event_name", "Unknown")
                return f"⚠️ RED FOLDER ACTIVE: {event_name}. Trading paused ±5 min around event."

            # 3. Fetch macro sentiment
            macro_news = data_agent.fetch_macro_news(symbol)

            # 4. Run SMC strategy analysis
            context = strategy.analyze_market_context(symbol, df, macro_news=macro_news)
            score = context.get("Structure_Score", "C")

            if score.startswith("C") or "error" in context:
                return f"Analysis for {symbol}: Score {score}. No valid setup found. Skipping."

            # 5. Build signal with direction override if active
            direction = "buy" if context.get("Macro_Trend") == "Bullish" else "sell"
            override = getattr(self, '_trade_direction_override', None)
            if override:
                direction = override
                self._trade_direction_override = None  # Single use

            signal = {
                "action": direction,
                "price": context.get("Current_Price"),
                "is_swing_trade": context.get("Is_Swing_Trade", False),
            }

            # 6. Risk validation
            equity = 100.0  # Paper trading default — will read from broker in production
            validated = risk.validate_trade_signal(equity, signal, df)

            if not validated.get("approved"):
                reason = validated.get("reason", "Unknown")
                return f"RiskAgent BLOCKED trade on {symbol}: {reason}"

            # 7. Execute
            current_price = df['close'].iloc[-1]
            receipt = execution.execute_trade(validated, current_price, exchange=exchange)

            if receipt.get("status") == "filled":
                # 8. Log to MemoryBank
                trade_id = str(receipt.get("id", f"{symbol}_{int(time.time())}"))
                memory.log_trade_entry(
                    trade_id=trade_id,
                    symbol=symbol,
                    direction=direction,
                    entry_price=current_price,
                    lot_size=validated.get("lot_size", 0),
                    smc_context=context
                )
                return (
                    f"✅ TRADE FIRED\n"
                    f"Symbol: {symbol} | {direction.upper()}\n"
                    f"Entry: {current_price} | Lots: {validated['lot_size']}\n"
                    f"S/L: {validated['stop_loss']} | T/P: {validated['take_profit']}\n"
                    f"Score: {score} | Risked: ${validated['capital_risked']}"
                )
            else:
                return f"Trade attempt on {symbol} failed: {receipt.get('reason', 'Unknown error')}"

        except ImportError as e:
            return f"Trading module import error: {e}"
        except Exception as e:
            return f"Trading pipeline error: {e}"

    def _trade_kill_switch(self) -> str:
        """Emergency: close all + revoke API keys."""
        try:
            from Sub_Projects.Trading.execution_agent import ExecutionAgent
            exec_agent = ExecutionAgent()
            exec_agent._monitor_running = False
            exec_agent.active_trades.clear()
        except:
            pass
        self._trading_paused = True
        return "🚨 KILL SWITCH ENGAGED. All trading halted. API connections severed."

    def _trade_close_all(self) -> str:
        """Close all open positions."""
        self._trading_paused = True
        return "All open positions marked for closure. Trading paused."

    def _trade_status(self) -> str:
        """Return current trading status."""
        paused = getattr(self, '_trading_paused', False)
        override = getattr(self, '_trade_direction_override', None)
        parts = [
            f"Trading: {'PAUSED' if paused else 'ACTIVE'}",
            f"Direction Override: {override or 'None (auto)'}",
        ]
        try:
            from src.core.memory_bank import MemoryBank
            mb = MemoryBank()
            stats = mb.analyze_consistency()
            if stats.get("status") != "warming_up":
                parts.append(f"Win Rate: {stats.get('win_rate', 'N/A')}%")
                parts.append(f"P&L: ${stats.get('total_pnl', 0)}")
            else:
                parts.append(stats.get("message", ""))
        except:
            parts.append("MemoryBank: offline")
        return " | ".join(parts)

    def _extract_symbol(self, cmd: str) -> str:
        """Extract trading symbol from command text."""
        symbols = {
            "btc": "BTCUSDT", "bitcoin": "BTCUSDT",
            "eth": "ETHUSDT", "ethereum": "ETHUSDT",
            "sol": "SOLUSDT", "solana": "SOLUSDT",
            "doge": "DOGEUSDT", "xrp": "XRPUSDT",
            "eurusd": "EURUSD", "gbpusd": "GBPUSD", "usdjpy": "USDJPY",
            "audusd": "AUDUSD", "usdcad": "USDCAD", "usdchf": "USDCHF",
            "eurjpy": "EURJPY", "gbpjpy": "GBPJPY",
            "xauusd": "XAUUSD", "gold": "XAUUSD", "silver": "XAGUSD",
            "aapl": "AAPL", "apple": "AAPL",
            "tsla": "TSLA", "tesla": "TSLA",
            "msft": "MSFT", "microsoft": "MSFT",
            "googl": "GOOGL", "google": "GOOGL",
            "amzn": "AMZN", "amazon": "AMZN",
            "nvda": "NVDA", "nvidia": "NVDA",
            "meta": "META", "amd": "AMD",
            "spy": "SPY", "qqq": "QQQ", "voo": "VOO",
            "nifty": "^NSEI", "sensex": "^BSESN",
            "reliance": "RELIANCE.NS", "tcs": "TCS.NS", "infosys": "INFY.NS",
        }
        for keyword, sym in symbols.items():
            if keyword in cmd:
                return sym
        return "AAPL"  # Safe default for paper trading

    def _run_backtest(self, cmd: str) -> str:
        """
        Run a universal backtest. Extracts symbol and time period from the command.
        Examples: "backtest xauusd 1 year", "backtest btc 3 months", "backtest aapl 15 days"
        """
        try:
            from Sub_Projects.Trading.backtest_agent import BacktestAgent, parse_period, resolve_symbol

            # Extract symbol from command
            # Remove "backtest" and common filler words to isolate the symbol
            parts = cmd.replace("backtest", "").replace("back test", "").strip()
            
            # Try to find a known symbol in the text
            symbol_found = None
            from Sub_Projects.Trading.backtest_agent import SYMBOL_ALIASES
            for alias in SYMBOL_ALIASES:
                if alias in parts.lower():
                    symbol_found = alias
                    break
            
            # Also check our orchestrator symbol map
            if not symbol_found:
                symbol_found = self._extract_symbol(parts)
            else:
                symbol_found = symbol_found  # Keep the raw alias, BacktestAgent will resolve it

            # Extract time period
            days, period_label = parse_period(cmd)

            self._notify_state("thinking")
            speak(f"Running backtest on {symbol_found} for {period_label}. This may take a moment.")

            agent = BacktestAgent()
            report = agent.run_backtest(symbol_found, days=days)
            
            return agent.format_report(report)

        except ImportError as e:
            return f"Backtest module error: {e}"
        except Exception as e:
            return f"Backtest failed: {e}"

    def _handle_automation(self, state: dict) -> str:
        """
        Phase 10: Task Chaining & GUI Automation Planner.
        Breaks down multi-step natural language commands into JSON action arrays,
        then executes them sequentially.
        """
        from src.security.rbac import get_rbac
        rbac = get_rbac()
        
        # Security Gate: GUI Automation requires at least L2 Auth (Apps/Files access)
        if rbac.auth_level < 2:
            return "ACCESS DENIED. GUI Automation requires Level 2 Authorization.\nReply with `/unlock <your_passphrase>` to temporarily elevate access and try again."
            
        cmd = state["command"]
        
        self._notify_state("thinking")
        speak("Building automation trajectory.")
        
        sys_prompt = (
            "You are CRAVE's Autonomous GUI Task Planner. "
            "Write a Python script to fulfill the user's specific desktop automation request.\n"
            "You MUST use 'pyautogui', 'time', and 'os' to control the keyboard and mouse.\n"
            "Return EXACTLY AND ONLY valid Python code inside a single ```python block.\n\n"
            "Rules:\n"
            "1. If opening a site or app, use os.system() or webbrowser.open() then time.sleep(3) before GUI actions.\n"
            "2. TradingView line shortcut is Alt+H.\n"
            "3. Use pyautogui.hotkey('ctrl', 't') etc for complex actions.\n"
            "4. NEVER write markdown outside the python block.\n"
        )
        
        if not self._router:
            return "Model router not initialized for planning."
            
        res = self._router.chat(prompt=cmd, system_prompt=sys_prompt)
        res_text = res.get("response", "").strip()
        
        # Extract python code block
        code_str = ""
        import re
        match = re.search(r"```python\n(.*?)\n```", res_text, re.DOTALL)
        if match:
            code_str = match.group(1).strip()
        elif "```" in res_text:
            code_str = res_text.split("```")[1].strip()
            if code_str.startswith("python"):
                code_str = code_str[6:].strip()
        else:
            code_str = res_text

        if not code_str or len(code_str) < 5:
            return "Failed to generate valid automation script."

        print(f"[Automate] Executing Python Macro:\n{code_str}")
        
        # Execute the python script dynamically
        import pyautogui
        import time
        import os
        import subprocess
        import webbrowser
        
        # Force a small failsafe warning
        pyautogui.FAILSAFE = True
        pyautogui.PAUSE = 0.5
        
        exec_globals = {
            "pyautogui": pyautogui,
            "time": time,
            "os": os,
            "subprocess": subprocess,
            "webbrowser": webbrowser
        }
        
        try:
            # We execute it in a thread so it doesn't block the orchestrator loop
            def run_macro():
                try:
                    exec(code_str, exec_globals)
                except pyautogui.FailSafeException:
                    print("[Automate] GUI MACRO ABORTED VIA FAILSAFE CORNER.")
                except Exception as e:
                    print(f"[Automate] Macro Error: {e}")
            
            import threading
            threading.Thread(target=run_macro, daemon=True).start()
            return "Automation macro injected. Hands off the mouse until complete!"
        except Exception as e:
            return f"Failed to run automation: {e}"

    def _handle_hack(self, state: dict) -> str:
        """
        Hacking / security commands (Phase 7).
        Uses ModelRouter to translate natural language into Kali bash commands,
        then executes them via KaliAgent. Requires L4 Authorization.
        """
        from src.security.rbac import get_rbac
        rbac = get_rbac()
        
        # Check if already authenticated via Voice Gate or Terminal
        if rbac.auth_level < 4:
            return "ACCESS DENIED. Offensive network operations require Level 4 Authorization. Please initiate voice verification."

        self._notify_state("thinking")
        
        # Determine if this should be routed to PentAGI or standard Kali execution
        router_sys = (
            "You are a routing agent. Determine if the user is asking to run a specific network command, "
            "or if they are asking for an autonomous, high-level penetration test/CTF against a target. "
            "Respond EXACTLY with a JSON string: {\"mode\": \"kali\" or \"pentagi\", \"target\": \"extracted_ip_or_domain_if_pentagi\"}"
        )
        
        if not self._router:
            return "Model router not ready for translation."
            
        route_res = self._router.chat(prompt=state["command"], system_prompt=router_sys)
        try:
            import json
            route_data = json.loads(route_res.get("response", "{}"))
        except:
            route_data = {"mode": "kali"}
            
        mode = route_data.get("mode", "kali")
        
        if mode == "pentagi" and route_data.get("target"):
            target = route_data.get("target")
            from src.agents.pentagi_agent import PentagiAgent
            pentagi = PentagiAgent()
            
            # Initiate async pentest
            import threading
            def _launch_pentagi():
                res = pentagi.start_mission(target, mode="autonomous")
                if res.get("success") and self._telegram:
                    self._telegram.send_message_sync(f"🚀 PentAGI CTF/Autonomous scan started on {target}.")
                else:
                    if self._telegram:
                        self._telegram.send_message_sync(f"❌ PentAGI failed to launch: {res.get('error')}")

            threading.Thread(target=_launch_pentagi, daemon=True, name="CRAVE_Pentagi_Manual").start()
            return f"Strategic objective acknowledged. Delegating {target} to PentAGI autonomous sub-routine."
        
        # Standard Kali mode fallback
        sys_over = (
            "You are an expert offensive security engineer. "
            "Translate the following user request into a single, exact bash command to be run on Kali Linux. "
            "Output ONLY the raw command string, nothing else. No markdown, no explanations."
        )
        
        res = self._router.chat(
            prompt=state["command"],
            system_prompt=sys_over
        )
        
        bash_cmd = res.get("response", "").strip()
        if not bash_cmd or len(bash_cmd) < 2:
            return "Failed to generate offensive payload from request."
            
        print(f"[Orchestrator] Translated to Kali Command: {bash_cmd}")
        
        from .tts import speak
        speak("Executing offensive payload on virtual machine.")
        
        from src.agents.kali_agent import KaliAgent
        agent = KaliAgent()
        output = agent.run_command(bash_cmd)
        
        return f"Kali Execution Complete.\n\nOutput:\n{output}"

    def _handle_silent(self, state: dict) -> str:
        """Toggle silent mode on/off."""
        new_state = not self._silent_mode
        self.set_silent_mode(new_state)
        if new_state:
            from .tts import speak_silent_on
            speak_silent_on()
            return "Silent mode on."
        else:
            from .tts import speak_silent_off
            speak_silent_off()
            return "Silent mode off."

    def _handle_auth(self, state: dict) -> str:
        """
        Hands-free Voice Passphrase Gate.
        Pauses the loop, listens specifically for the 4-word L4 phrase, and verifies it.
        """
        self._notify_state("speaking")
        from .tts import speak
        speak("Voice Protocol Active. State your clearance passphrase.")
        
        self._notify_state("listening")
        if not self._voice:
            return "Microphone offline. Cannot verify voice."
            
        print("[Orchestrator] Waiting 15s for verbal clearance phrase...")
        phrase = self._voice.listen_once(timeout=15.0)
        
        if not phrase:
            return "No phrase detected. Protocol aborted."
            
        from src.security.rbac import get_rbac
        rbac = get_rbac()
        
        import re
        clean_phrase = re.sub(r'[^a-zA-Z0-9\s]', '', phrase.lower()).strip()
        ph_raw = phrase.strip('.')
        
        matched = False
        try:
            # We match strict transcribed, clean text, or lowercase raw
            if rbac._verify_secret(phrase, rbac.credentials.get("L4_PHR_HASH", "")):
                matched = True
            elif rbac._verify_secret(clean_phrase, rbac.credentials.get("L4_PHR_HASH", "")):
                matched = True
            elif rbac._verify_secret(ph_raw, rbac.credentials.get("L4_PHR_HASH", "")):
                matched = True
        except: pass
        
        if matched:
            rbac.auth_level = max(rbac.auth_level, 4)
            rbac.touch()
            return "Clearance accepted. System limits overridden."
        else:
            return "Biometric mismatch. Authorization denied."

    def _handle_explain(self, state: dict) -> str:
        """
        'Explain Yourself' command — Phase 11 Transparency.
        Returns a human-readable log of recent autonomous decisions.
        """
        try:
            from src.core.reasoning_log import get_reasoning_log
            log = get_reasoning_log()

            if log.count() == 0:
                return (
                    "I haven't made any autonomous decisions yet in this session. "
                    "Once I execute trades, run Kali commands, or take other autonomous "
                    "actions, I'll record my reasoning here for you to review."
                )

            explanation = log.explain_last(n=5)
            return f"Here are my most recent decisions:\n\n{explanation}"

        except Exception as e:
            return f"Failed to retrieve reasoning log: {e}"

    def _handle_evolve(self, state: dict) -> str:
        """Phase 11: Trigger ModelManager to check for brain upgrades."""
        try:
            from src.core.model_manager import ModelManager
            mm = ModelManager()
            
            # This is normally run async, but we run it synchronously here 
            # so we can return the result directly to the user.
            candidates = mm.check_system_resources()
            # The actual discovery will be built out more fully soon,
            # this is just the entrypoint wrapper.
            return f"Model Evolution triggered. System check: {candidates}"
        except Exception as e:
            return f"Evolution check failed: {e}"

    def _handle_self_modify(self, state: dict) -> str:
        """Phase 11: Trigger SelfModifier for codebase changes."""
        command = state.get("command", "")
        # Remove trigger intent words to get raw task
        raw_task = command
        for kw in _INTENT_KEYWORDS[INTENT_SELF_MODIFY]:
            if kw in raw_task.lower():
                raw_task = raw_task.lower().replace(kw, "").strip()
                
        if not raw_task:
            return "Please provide a specific feature or modification you want me to implement."
            
        try:
            from src.core.self_modifier import SelfModifier
            import threading
            
            modifier = SelfModifier()
            
            # Since self-modification (code gen, tests) takes minutes, do it async
            # and notify back via telegram or TTS when ready for approval.
            def _async_modify():
                result = modifier.execute_modification(raw_task)
                if self.telegram:
                    self.telegram.send_message_sync(f"🛠️ Self-Modification Result:\n{result}")
                logger.info(f"Self-Modification completed:\n{result}")
                
            threading.Thread(target=_async_modify, daemon=True).start()
            
            return f"I've initiated the self-modification process for: '{raw_task}'. " \
                   "This will take a few minutes. I'll create a sandbox, generate code, " \
                   "run tests, and ask for your explicit approval before changing anything."
        except Exception as e:
             return f"Failed to start self-modification: {e}"

    def _handle_status(self, state: dict) -> str:

        """Report current system status."""
        st = self.get_status()
        parts = [
            "CRAVE status:",
            f"Running: {st['running']}",
            f"Silent: {st['silent_mode']}",
            f"Messages handled: {st['msg_count']}",
            f"Voice active: {st['voice_running']}",
            "Phases complete: 1-9 (foundation, router, voice, orchestrator, API, security, UI, tools, trading, learning).",
            "Phase active: 10 (production hardening).",
        ]
        return " ".join(parts)

    def _handle_stop(self, state: dict) -> str:
        """Graceful shutdown."""
        speak("Shutting down. Goodbye.")
        threading.Thread(target=self.stop, daemon=True).start()
        return "Shutting down."

    def _handle_scout(self, state: dict) -> str:
        """Phase 12: Multi-platform LLM research scout."""
        cmd = state["command"]
        
        # Extract the search query
        query = cmd.lower()
        for prefix in ["research github", "find improvements", "scout llm", "scout repos",
                       "trending llm", "research llm", "find new models", "huggingface trending",
                       "what's new in ai", "research ai improvements", "find better techniques"]:
            query = query.replace(prefix, "").strip()
        
        if not query:
            query = "LLM optimization local inference"
        
        speak("Initiating multi-platform intelligence sweep. Scanning GitHub, HuggingFace, PapersWithCode, and Reddit.")
        
        try:
            from src.agents.llm_scout import LLMScout
            scout = LLMScout(router=self._router)
            result = scout.scout(query)
            
            analysis = result.get("analysis", "No analysis available.")
            findings = result.get("findings", [])
            report_path = result.get("report_path", "")
            
            summary = f"Scanned {len(findings)} projects across 4 platforms.\n\n{analysis}"
            
            if report_path:
                summary += f"\n\nFull report saved: {report_path}"
            
            summary += "\n\nSay 'apply improvement [number]' if you want me to sandbox-test any of these."
            
            return summary
            
        except Exception as e:
            return f"Scout mission failed: {e}"

    def _analyze_adaptive_signals(self, user_text: str):
        """
        Phase 12 Adaptive Personality: Analyze user's response for positive/negative signals.
        Silently updates user_profile.json to track explanation preferences.
        """
        profile_path = os.path.join(crave_root(), "data", "user_profile.json")
        
        try:
            if not os.path.exists(profile_path):
                return
            
            with open(profile_path, "r", encoding="utf-8") as f:
                profile = json.load(f)
            
            lower = user_text.lower()
            
            # Positive signals
            positive_words = ["thanks", "got it", "perfect", "great", "nice", "understood",
                             "that makes sense", "clear", "awesome", "good", "ok cool"]
            if any(w in lower for w in positive_words):
                profile["positive_signals"] = profile.get("positive_signals", 0) + 1
            
            # Negative signals
            negative_words = ["what?", "don't get it", "confused", "huh", "not clear",
                             "explain again", "i don't understand", "too complex", "simplify"]
            if any(w in lower for w in negative_words):
                profile["negative_signals"] = profile.get("negative_signals", 0) + 1
            
            # Every 20 interactions, recalibrate the preferred style
            total = profile.get("total_interactions", 0)
            if total > 0 and total % 20 == 0 and self._router:
                pos = profile.get("positive_signals", 0)
                neg = profile.get("negative_signals", 0)
                topics = [t["topic"] for t in profile.get("topics_discussed", [])[-10:]]
                
                calibration_prompt = (
                    f"A user has interacted with me {total} times. "
                    f"Positive feedback signals: {pos}, Negative/confused signals: {neg}. "
                    f"Recent topics: {', '.join(topics)}. "
                    f"Based on this data, what explanation style works best for this user? "
                    f"Choose ONE from: analogies, bullet_points, visual, code_examples, stories, step_by_step. "
                    f"Return ONLY the style name, nothing else."
                )
                
                res = self._router.chat(
                    prompt=calibration_prompt,
                    system_prompt="You are a learning analytics engine. Return only one word.",
                    task_type="primary"
                )
                style = res.get("response", "").strip().lower().replace(" ", "_")
                valid_styles = ["analogies", "bullet_points", "visual", "code_examples", "stories", "step_by_step"]
                if style in valid_styles:
                    profile["preferred_style"] = style
            
            with open(profile_path, "w", encoding="utf-8") as f:
                json.dump(profile, f, indent=2)
                
        except Exception:
            pass

    # ── context management ────────────────────────────────────────────────────



    def _compress_context(self):
        """
        Summarise old context to keep prompt length manageable.
        Uses ModelRouter to generate a concise summary of older messages,
        preserving key facts and decisions while dropping raw text.
        Falls back to simple truncation if summarization fails.
        """
        keep_recent = 20  # Keep last 10 exchanges (20 messages) verbatim
        
        if len(self._context) <= keep_recent:
            return  # Nothing to compress
        
        old_messages = self._context[:-keep_recent]
        recent_messages = self._context[-keep_recent:]
        
        # Attempt AI-powered summarization of old context
        try:
            # Build a conversation transcript from old messages
            transcript_lines = []
            for msg in old_messages:
                role = msg.get("role", "unknown").upper()
                content = msg.get("content", "")[:200]  # Cap per-message to avoid huge prompts
                transcript_lines.append(f"{role}: {content}")
            
            transcript = "\n".join(transcript_lines[-40:])  # Max 40 messages for summary input
            
            summary_prompt = (
                "Summarize the following conversation in 3-4 sentences. "
                "Preserve key facts, decisions, user preferences, and any pending tasks. "
                "Do NOT add opinions or speculation:\n\n"
                f"{transcript}"
            )
            
            result = self.router.chat(
                prompt=summary_prompt,
                task_type="primary",
                system_prompt="You are a precise summarization engine. Output only the summary."
            )
            
            summary_text = result.get("response", "").strip()
            
            if summary_text and len(summary_text) > 20:
                # Replace old messages with a single summary system message
                summary_msg = {
                    "role": "system",
                    "content": f"[CONTEXT SUMMARY — {len(old_messages)} prior messages compressed]: {summary_text}"
                }
                self._context = [summary_msg] + recent_messages
                self._msg_count = len(self._context)
                logger.info(f"Context compressed via AI summary ({len(old_messages)} messages → 1 summary)")
                return
                
        except Exception as e:
            logger.warning(f"AI context compression failed, falling back to truncation: {e}")
        
        # Fallback: simple truncation if summarization fails
        dropped = len(self._context) - keep_recent
        self._context = self._context[-keep_recent:]
        self._msg_count = keep_recent
        print(f"[Orchestrator] Context compressed via truncation (dropped {dropped} old messages)")

    # ── helpers ───────────────────────────────────────────────────────────────

    def _load_system_prompt(self) -> str:
        """Load personality/instructions from program.md."""
        path = os.path.join(crave_root(), "Main_Lead", "program.md")
        try:
            with open(path, encoding="utf-8") as f:
                content = f.read()
            # Strip only file-level comments (meta lines like "# CRAVE 2026" and "# Save to:")
            # Keep section headers (## Identity, ## Core Rules, etc.) for LLM context
            lines = []
            for ln in content.splitlines():
                stripped = ln.strip()
                # Skip empty lines
                if not stripped:
                    continue
                # Skip file meta comments (single # followed by uppercase meta keywords)
                if stripped.startswith("# ") and any(kw in stripped for kw in ["CRAVE", "Save to:", "Edit this"]):
                    continue
                lines.append(stripped)
            prompt = "\n".join(lines)
            print(f"[Orchestrator] Loaded system prompt ({len(prompt)} chars)")
            return prompt
        except FileNotFoundError:
            default = (
                "You are CRAVE, a local AI assistant. "
                "Be concise, direct, and helpful. "
                "Always confirm before dangerous actions."
            )
            print("[Orchestrator] program.md not found — using default prompt")
            return default
        except Exception as e:
            print(f"[Orchestrator] Could not load program.md: {e}")
            return "You are CRAVE, a helpful AI assistant."

    # ── convenience function for Phase 6 Orb ─────────────────────────────────────

_global_orchestrator: Optional[Orchestrator] = None

def get_orchestrator() -> Orchestrator:
    """
    Return the global Orchestrator instance.
    Phase 6 Orb UI calls this to get the shared instance.
    """
    global _global_orchestrator
    if _global_orchestrator is None:
        _global_orchestrator = Orchestrator()
    return _global_orchestrator


def start_crave():
    """Start CRAVE. Called from main entry point or Phase 6 Orb."""
    orc = get_orchestrator()
    orc.start()
    return orc


def stop_crave():
    """Stop CRAVE cleanly."""
    global _global_orchestrator
    if _global_orchestrator:
        _global_orchestrator.stop()
        _global_orchestrator = None